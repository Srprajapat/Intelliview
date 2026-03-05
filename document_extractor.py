import os
import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    with fitz.open(file_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file (paragraphs + tables)."""
    doc = Document(docx_path)
    full_text = []

    # 1. Extract Paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 2. Extract Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_text))

    return '\n'.join(full_text)

def extract_text_auto(file_path):
    """
    Automatically detects file type and extracts text.
    Returns a dictionary with status and data.
    """
    # 1. Check if file exists
    if not os.path.exists(file_path):
        return {
            "text": None,
            "text_extraction": False,
            "message": f"File not found: {file_path}"
        }

    # 2. Get the extension (lowercase for comparison)
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()

    try:
        extracted_text = ""
        
        # 3. Route to the correct function
        if extension == ".pdf":
            extracted_text = extract_text_from_pdf(file_path)
        
        elif extension == ".docx":
            extracted_text = extract_text_from_docx(file_path)
        
        else:
            # File type not supported
            return {
                "text": None,
                "text_extraction": False,
                "message": f"Text extraction unsuccessful: Unsupported file type '{extension}'"
            }

        # 4. Success Case
        return {
            "text": extracted_text,
            "text_extraction": True,
            "message": "Text extraction successful"
        }

    except Exception as e:
        # 5. Error Case (Corrupted file, password protected, etc.)
        return {
            "text": None,
            "text_extraction": False,
            "message": f"Text extraction unsuccessful: {str(e)}"
        }

# --- Example Usage ---
# result = extract_text_auto("resume.pdf")
# print(result)